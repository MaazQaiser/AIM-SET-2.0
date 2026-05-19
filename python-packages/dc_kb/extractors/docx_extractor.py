from __future__ import annotations

from pathlib import Path

from dc_kb.models import ExtractedDocument, TextChunk


def extract_docx(path: Path) -> ExtractedDocument:
    from docx import Document

    doc = Document(str(path))
    parts = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    text = "\n\n".join(parts)
    chunks = [TextChunk(text=text, metadata={"source": path.name})] if text else []
    return ExtractedDocument(chunks=chunks, metadata={"format": "docx"})
